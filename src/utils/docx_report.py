"""DOCX report generator for accessibility scan results."""
import io
from datetime import datetime
from typing import Dict, List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.models import ScanResult, PageResult, AccessibilityViolation


class DOCXReportGenerator:
    """Generate DOCX reports for accessibility scans."""

    def __init__(self, scan_result: ScanResult):
        self.scan_result = scan_result
        self.document = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Set up document styles for consistency."""
        # Set default font
        style = self.document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

    def generate(self) -> io.BytesIO:
        """
        Generate the complete DOCX report.
        
        Returns:
            BytesIO object containing the DOCX file
        """
        self._add_title()
        self._add_summary()
        self._add_scan_details()
        self._add_violations_by_severity()
        self._add_pages_with_violations()

        # Save to BytesIO
        docx_file = io.BytesIO()
        self.document.save(docx_file)
        docx_file.seek(0)
        return docx_file

    def _add_title(self):
        """Add report title."""
        title = self.document.add_heading('Accessibility Scan Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add scan date
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(
            f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        )
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)

        self.document.add_paragraph()  # Spacing

    def _add_summary(self):
        """Add executive summary section."""
        self.document.add_heading('Executive Summary', 1)

        summary_data = [
            ('Website URL', self.scan_result.start_url),
            ('Total Pages Scanned', str(self.scan_result.pages_scanned)),
            ('Pages with Violations', str(self.scan_result.pages_with_violations)),
            ('Total Violations Found', str(self.scan_result.total_violations)),
            ('Scan Status', self.scan_result.status.upper()),
            ('Scan Date', self.scan_result.start_time.strftime('%B %d, %Y')),
        ]

        if self.scan_result.duration:
            minutes = int(self.scan_result.duration // 60)
            seconds = int(self.scan_result.duration % 60)
            summary_data.append(('Scan Duration', f"{minutes}m {seconds}s"))

        # Create summary table
        table = self.document.add_table(rows=len(summary_data), cols=2)
        table.style = 'Light Grid Accent 1'

        for i, (label, value) in enumerate(summary_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # Bold the labels
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        self.document.add_paragraph()  # Spacing

    def _add_scan_details(self):
        """Add scan configuration details."""
        self.document.add_heading('Scan Configuration', 1)

        config_para = self.document.add_paragraph()
        config_items = [
            f"Maximum Pages: {self.scan_result.max_pages}",
            f"Maximum Depth: {self.scan_result.max_depth}",
            f"Same Domain Only: {'Yes' if self.scan_result.same_domain_only else 'No'}",
            f"Scan Mode: {self.scan_result.scan_mode.upper()}"
        ]
        config_para.add_run('\n'.join(config_items))

        self.document.add_paragraph()  # Spacing

    def _add_violations_by_severity(self):
        """Add violations breakdown by severity."""
        self.document.add_heading('Violations by Severity', 1)

        severity_counts = self.scan_result.get_violations_by_severity()

        # Create table
        table = self.document.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Severity Level'
        header_cells[1].text = 'Count'
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True

        # Data rows
        severities = [
            ('Error', severity_counts.get('error', 0)),
            ('Warning', severity_counts.get('warning', 0)),
            ('Alert', severity_counts.get('alert', 0))
        ]

        for i, (severity, count) in enumerate(severities, start=1):
            row = table.rows[i]
            row.cells[0].text = severity
            row.cells[1].text = str(count)

        self.document.add_paragraph()  # Spacing

    def _add_pages_with_violations(self):
        """Add detailed violations grouped by page and severity."""
        self.document.add_heading('Detailed Violations by Page', 1)

        # Group pages by whether they have violations
        pages_with_issues = [p for p in self.scan_result.page_results if p.has_violations]

        if not pages_with_issues:
            self.document.add_paragraph('No violations found. Excellent work!')
            return

        for page in pages_with_issues:
            self._add_page_violations(page)

    def _add_page_violations(self, page: PageResult):
        """Add violations for a single page."""
        # Page heading
        page_heading = self.document.add_heading(level=2)
        page_heading.add_run(page.url)

        # Page info
        info_para = self.document.add_paragraph()
        info_run = info_para.add_run(
            f"Total Violations: {page.violation_count}"
        )
        info_run.font.italic = True
        info_run.font.size = Pt(10)

        # Group violations by severity
        violations_by_severity = self._group_violations_by_severity(page.violations)

        # Add each severity section
        for severity in ['error', 'warning', 'alert']:
            violations = violations_by_severity.get(severity, [])
            if violations:
                self._add_severity_section(severity, violations)

        self.document.add_paragraph()  # Spacing between pages

    def _group_violations_by_severity(
        self, 
        violations: List[AccessibilityViolation]
    ) -> Dict[str, List[AccessibilityViolation]]:
        """Group violations by their effective severity."""
        grouped = {'error': [], 'warning': [], 'alert': []}
        
        for violation in violations:
            severity = violation.effective_severity
            if severity in grouped:
                grouped[severity].append(violation)
        
        return grouped

    def _add_severity_section(self, severity: str, violations: List[AccessibilityViolation]):
        """Add a section for violations of a specific severity."""
        # Severity subheading
        severity_heading = self.document.add_heading(level=3)
        severity_label = severity.upper()
        severity_heading.add_run(f"{severity_label}S ({len(violations)})")

        # Add each violation
        for i, violation in enumerate(violations, start=1):
            self._add_violation_details(violation, i)

    def _add_violation_details(self, violation: AccessibilityViolation, number: int):
        """Add details for a single violation."""
        # Violation title
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run(f"{number}. {violation.description}")
        title_run.font.bold = True
        title_run.font.size = Pt(11)

        # Help text
        if violation.help:
            help_para = self.document.add_paragraph(style='List Bullet')
            help_para.add_run(f"How to Fix: {violation.help}")

        # WCAG tags
        if violation.tags:
            wcag_tags = [tag for tag in violation.tags if 'wcag' in tag.lower()]
            if wcag_tags:
                tags_para = self.document.add_paragraph(style='List Bullet')
                tags_run = tags_para.add_run(f"WCAG Guidelines: {', '.join(wcag_tags)}")
                tags_run.font.size = Pt(10)

        # Affected elements count
        if violation.nodes:
            nodes_para = self.document.add_paragraph(style='List Bullet')
            nodes_run = nodes_para.add_run(f"Affected Elements: {len(violation.nodes)}")
            nodes_run.font.size = Pt(10)
            nodes_run.font.color.rgb = RGBColor(128, 128, 128)

        # More info link
        if violation.help_url:
            link_para = self.document.add_paragraph(style='List Bullet')
            link_run = link_para.add_run(f"More Information: {violation.help_url}")
            link_run.font.size = Pt(9)
            link_run.font.color.rgb = RGBColor(0, 0, 255)
            link_run.font.underline = True

        self.document.add_paragraph()  # Spacing between violations


def generate_docx_report(scan_result: ScanResult) -> io.BytesIO:
    """
    Generate a DOCX report for the given scan result.
    
    Args:
        scan_result: The scan result to generate a report for
        
    Returns:
        BytesIO object containing the DOCX file
    """
    generator = DOCXReportGenerator(scan_result)
    return generator.generate()

